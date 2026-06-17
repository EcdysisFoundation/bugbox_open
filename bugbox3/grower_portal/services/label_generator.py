import logging
import os
import re
import secrets
from copy import deepcopy
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import _Cell, _Row

from bugbox3.core.models import PrivateSiteContent

from ..constants import (
    AVALANCHE_REFRIGERATED_SAMPLES,
    AVALANCHE_ROOM_TEMP_SAMPLES,
    LABEL_IGNITE_TEMPLATE_SLUG,
    LABEL_OUTER_TEMPLATE_SLUG,
    LABEL_TEMPLATE_SLUG,
    SAMPLE_TYPES,
)
from ..models import SampleCode, SiteTransect

logger = logging.getLogger(__name__)


def _label_fill_debug_enabled() -> bool:
    """per-cell label fill logs"""
    return os.environ.get('BUGBOX_LABEL_FILL_DEBUG', '').strip().lower() in (
        '1',
        'true',
        'yes',
        'on',
    )


def _label_fill_info_enabled() -> bool:
    """fill summaries"""
    if _label_fill_debug_enabled():
        return True
    return os.environ.get('BUGBOX_LABEL_FILL_INFO', '').strip().lower() in (
        '1',
        'true',
        'yes',
        'on',
    )


class LabelGenerator:
    """Service for generating label documents"""

    _MIN_LABEL_SLOT_CELL_WIDTH_TWIPS = 1000

    def __init__(
            self,
            project_type,
            cluster_number,
            year,
            sample_types,
            labels_per_type,
            created_by,
            label_category='inner'):
        self.project_type = project_type
        self.cluster_number = cluster_number
        self.year = year
        self.sample_types = sample_types
        self.labels_per_type = labels_per_type
        self.created_by = created_by
        self.label_category = label_category
        self.generated_codes = []

    def get_sample_type_display(self, sample_type_code):
        choices = dict(SAMPLE_TYPES)
        return choices.get(sample_type_code, sample_type_code)

    def generate_unique_sample_codes(self, count):
        """
        Generate unique sample codes.

        - Avalanche: example (cluster=82): 821231
        - Ignite: numeric sequential (Ignite quick-generate uses generate_unique_site_codes).
        """
        if self.project_type == 'avalanche':
            return self._generate_unique_avalanche_sample_codes(count)

        codes = []
        all_codes = SampleCode.objects.filter(project_type='ignite').values_list('code', flat=True)

        max_code = 0
        for code in all_codes:
            try:
                num = int(code)
                max_code = max(max_code, num)
            except ValueError:
                digits = re.findall(r'\d+', code)
                if digits:
                    num = int(digits[-1])
                    max_code = max(max_code, num)

        for i in range(count):
            new_code = str(max_code + i + 1).zfill(6)
            codes.append(new_code)

        return codes

    def _generate_unique_avalanche_sample_codes(self, count):
        """
        Avalanche sample code format: <cluster_number><4 random digits>

        - cluster_number comes from the Quick Generate form
        """
        cluster_digits = re.sub(r'\D+', '', str(self.cluster_number or '')).strip()
        if not cluster_digits:
            raise ValueError('Cluster number must contain digits to generate Avalanche sample codes.')

        if count > 10_000:
            raise ValueError(
                f'Cannot generate {count} Avalanche codes with only 4 digits after the cluster '
                f'number (max 10000 per cluster).'
            )

        generated = []
        generated_set = set()

        max_attempts = max(50_000, count * 200)
        attempts = 0

        while len(generated) < count:
            attempts += 1
            if attempts > max_attempts:
                raise ValueError(
                    f'Unable to generate {count} unique Avalanche sample codes for cluster '
                    f'{cluster_digits} after {max_attempts} attempts. The 4-digit namespace may '
                    f'be close to exhausted.'
                )

            suffix = secrets.randbelow(10_000)
            candidate = f"{cluster_digits}{suffix:04d}"

            if candidate in generated_set:
                continue

            if SampleCode.objects.filter(code=candidate).exists():
                continue

            generated.append(candidate)
            generated_set.add(candidate)

        return generated

    def generate_unique_site_codes(self, count):
        """Generate Ignite site codes with auto-incrementing logic"""
        year_prefix = str((self.year % 10) - 1)

        # Find last site code for this year
        last_site = SampleCode.objects.filter(
            project_type='ignite',
            year=self.year,
            site_code_numeric__isnull=False
        ).order_by('-site_code_numeric').first()

        if last_site:
            last_num = last_site.site_code_numeric
            cluster_offset = ((last_num % 1000) - 1) // 30
            next_cluster_start = int(year_prefix + '000') + ((cluster_offset + 1) * 30) + 1
        else:
            # First cluster for this year starts at X001
            next_cluster_start = int(year_prefix + '001')

        codes = []
        for i in range(count):
            site_num = next_cluster_start + i
            code = str(site_num)
            codes.append(code)

        return codes, next_cluster_start

    def save_sample_codes(self, codes, project_type='avalanche'):
        sample_code_objects = [
            SampleCode(
                code=code,
                project_type=project_type,
                cluster_number=self.cluster_number,
                year=self.year,
                site_code_numeric=int(code) if project_type == 'ignite' else None,
                is_active=True,
                is_used=False,
                created_by=self.created_by
            )
            for code in codes
        ]
        SampleCode.objects.bulk_create(sample_code_objects)
        self.generated_codes = codes

    def create_label_text(self, sample_type_display, transect_code):
        """Create label text based on project type"""
        if self.project_type == 'avalanche':
            # Avalanche format: Sample Type / Cluster-Year / Transect Code
            return f"{sample_type_display}\nCluster {self.cluster_number} – {self.year}\n{transect_code}"
        else:
            # ignite (1000 farms) format: Cluster-Year / Sample Type / Transect Code
            return f"Cluster {self.cluster_number} – {self.year}\n{sample_type_display}\n{transect_code}"

    def _load_template(self):
        """Load the template document from PrivateSiteContent"""
        if self.project_type == 'ignite':
            template_slug = LABEL_IGNITE_TEMPLATE_SLUG
        else:
            template_slug = LABEL_OUTER_TEMPLATE_SLUG if self.label_category == 'outer' else LABEL_TEMPLATE_SLUG

        try:
            template_content = PrivateSiteContent.objects.get(title=template_slug)

            if not template_content.file:
                raise ValueError(
                    f"Label template '{template_slug}' found but file is missing. Please upload it in Django Admin"
                )

            with template_content.file.open('rb') as file:
                file_content = file.read()
                file_buffer = BytesIO(file_content)

                file_buffer.seek(0)
                header = file_buffer.read(8)
                file_buffer.seek(0)

                if header[:4] == b'\xd0\xcf\x11\xe0' or header[:2] != b'PK':
                    raise ValueError(
                        f"Template file '{template_slug}' is in '.doc' format. "
                        f"Please convert it to '.docx' format and re-upload. "
                    )

                try:
                    return Document(file_buffer)
                except Exception as e:
                    error_msg = str(e).lower()
                    if (
                        'not a word file' in error_msg or
                        'content type' in error_msg or
                        'not a zip file' in error_msg or
                        'bad zipfile' in error_msg
                    ):
                        raise ValueError(
                            f"Template file '{template_slug}' could not be "
                            f"loaded. The file must be in '.docx' format. "
                            f"Original error: {str(e)}"
                        ) from e
                    else:
                        raise

        except PrivateSiteContent.DoesNotExist:
            project_label = "Ignite" if self.project_type == 'ignite' else "Avalanche"
            raise ValueError(
                f"Label template '{template_slug}' not found in Private Site Content. "
                f"Please upload the {project_label} template file in Django Admin."
            )
        except Exception as e:
            raise ValueError(
                f"Error loading label template: {str(e)}"
            ) from e

    def _replace_cell_text_preserving_format(self, cell, new_text):
        """Replace cell text by modifying existing runs"""
        if not cell.paragraphs:
            cell.add_paragraph()

        paragraph = cell.paragraphs[0]

        if paragraph.runs:
            runs = list(paragraph.runs)

            first_run = runs[0]
            first_run.text = new_text

            for r in runs[1:]:
                r._element.getparent().remove(r._element)
        else:
            if paragraph.text.strip():
                paragraph.clear()

            run = paragraph.add_run(new_text)
            run.font.size = Pt(11)
            run.font.bold = True

        return paragraph

    def _remove_pages_after_last_filled_cell(self, doc):
        """Find the last filled cell (with sample type and/or transect code) and remove all pages after it"""
        if not doc.tables:
            return

        last_filled_table_idx = -1

        for table_idx in range(len(doc.tables) - 1, -1, -1):
            try:
                table = doc.tables[table_idx]
                for row_idx, row in enumerate(table.rows):
                    try:
                        if not hasattr(row, 'cells'):
                            continue
                        for col_idx, cell in enumerate(row.cells):
                            try:
                                cell_text = cell.text.strip()
                                if cell_text and cell_text.upper() != 'LABEL' and len(cell_text) > 1:
                                    has_sample_type = any(char.isalpha() for char in cell_text)
                                    has_transect_code = any(char.isdigit() for char in cell_text)
                                    has_multiline = '\n' in cell_text

                                    if (has_sample_type or has_transect_code) and (has_transect_code or has_multiline):
                                        if last_filled_table_idx < 0 or table_idx > last_filled_table_idx:
                                            last_filled_table_idx = table_idx
                            except (AttributeError, IndexError):
                                continue
                    except (IndexError, AttributeError):
                        continue
            except (IndexError, AttributeError):
                continue

        if last_filled_table_idx >= 0 and last_filled_table_idx < len(doc.tables) - 1:
            tables_to_remove = list(range(len(doc.tables) - 1, last_filled_table_idx, -1))

            for idx in tables_to_remove:
                try:
                    if idx < len(doc.tables):
                        table_element = doc.tables[idx]._tbl
                        parent = table_element.getparent()
                        if parent is not None:
                            parent.remove(table_element)
                except (IndexError, AttributeError, ValueError):
                    pass

        try:
            if doc.tables:
                last_table_idx = len(doc.tables) - 1
                if last_table_idx >= 0:
                    last_table_element = doc.tables[last_table_idx]._tbl
                    parent = last_table_element.getparent()
                    if parent is not None:
                        try:
                            last_table_index = list(parent).index(last_table_element)
                        except ValueError:
                            last_table_index = -1

                        if last_table_index >= 0:
                            all_elements = list(parent)
                            elements_after_table = all_elements[last_table_index + 1:]

                            for elem in reversed(elements_after_table):
                                try:
                                    if elem.tag.endswith('}p'):
                                        para_text = ''.join(elem.itertext()).strip()
                                        if not para_text:
                                            try:
                                                parent.remove(elem)
                                            except (ValueError, AttributeError):
                                                pass
                                except (ValueError, AttributeError):
                                    pass
        except (IndexError, AttributeError, ValueError):
            pass

    def _template_has_placeholders(self, doc):
        """Check if template has LABEL or {{label}} placeholders"""
        if not doc.tables:
            return False

        for table in doc.tables:
            try:
                for row in table.rows:
                    try:
                        if not hasattr(row, 'cells'):
                            continue
                        for cell in row.cells:
                            try:
                                cell_text_upper = cell.text.upper()
                                # Check for LABEL placeholders
                                if ('LABEL' in cell_text_upper or
                                    '{{LABEL}}' in cell_text_upper or
                                        '{LABEL}' in cell_text_upper):
                                    return True
                            except (AttributeError, IndexError):
                                continue
                    except (IndexError, AttributeError):
                        continue
            except (IndexError, AttributeError):
                continue
        return False

    def _distribute_labels_to_columns(self, doc, all_labels):
        """Distribute labels across template columns"""
        if not doc.tables or not all_labels:
            return [all_labels]

        num_columns = 0
        valid_table_found = False

        for table_idx, table in enumerate(doc.tables):
            if not table.rows:
                continue

            try:
                first_row = table.rows[0]
                cells = list(first_row.cells)  # This will fail if merged cells cause issues

                print(f"DEBUG: Table {table_idx} - First row has {len(cells)} cells")
                for idx, cell in enumerate(cells):
                    cell_text = cell.text
                    cell_text_upper = cell_text.upper()
                    print(f"DEBUG: Table {table_idx} Cell {idx} text: '{cell_text}'")
                    # Check for LABEL, {{label}}, {label}, or label placeholder
                    if ('LABEL' in cell_text_upper or
                        '{{LABEL}}' in cell_text_upper or
                            '{LABEL}' in cell_text_upper):
                        num_columns += 1
                        print(f"DEBUG: Found LABEL placeholder in table {table_idx} cell {idx}")

                # If we found LABEL columns in this table, stop looking
                if num_columns > 0:
                    valid_table_found = True
                    print(f"DEBUG: Using table {table_idx} as template (found {num_columns} LABEL columns)")
                    break

            except (IndexError, AttributeError, ValueError) as e:
                print(f"DEBUG: Skipping table {table_idx} due to error: {e}")
                continue

        print(f"DEBUG: Found {num_columns} LABEL columns in template (valid_table_found={valid_table_found})")

        # If no columns found or only one column, return all labels in single column
        if num_columns <= 1:
            return [all_labels]

        # Distribute labels across columns
        labels_by_column = [[] for _ in range(num_columns)]
        for idx, label in enumerate(all_labels):
            col_idx = idx % num_columns
            labels_by_column[col_idx].append(label)

        return labels_by_column

    def _fill_template_with_placeholders(self, doc, labels_by_column):
        """Fill template using placeholder replacement"""
        if not doc.tables:
            return doc

        if not labels_by_column:
            return doc

        num_sample_types = len(labels_by_column)
        label_index_by_sample_type = [0] * num_sample_types

        if not doc.tables or len(doc.tables) == 0:
            return doc

        first_table = None
        for table in doc.tables:
            try:
                if len(table.rows) > 0:
                    first_row = table.rows[0]
                    if hasattr(first_row, 'cells') and len(first_row.cells) > 0:
                        _ = first_row.cells[0]
                        first_table = table
                        break
            except (IndexError, AttributeError):
                continue

        if first_table is None:
            return doc

        accessible_template_table = None
        for table in doc.tables:
            try:
                if table.rows and len(table.rows) > 0:
                    _ = list(table.rows[0].cells)  # Test if cells are accessible
                    accessible_template_table = table
                    break
            except (IndexError, AttributeError, ValueError):
                continue

        if accessible_template_table is None:
            accessible_template_table = first_table

        original_template_table_xml = deepcopy(accessible_template_table._tbl)
        original_table_style = first_table.style if first_table.style else None

        original_style_val = None
        try:
            original_tbl_pr = first_table._tbl.tblPr
            if original_tbl_pr is not None:
                for child in original_tbl_pr:
                    if 'tblStyle' in child.tag:
                        original_style_val = child.get(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                        break
        except Exception:
            pass

        template_cols_with_label = []
        working_table = None
        for table in doc.tables:
            try:
                if table.rows and len(table.rows) > 0:
                    _ = list(table.rows[0].cells)
                    working_table = table
                    break
            except (IndexError, AttributeError, ValueError):
                continue

        if working_table and working_table.rows:
            try:
                first_row = working_table.rows[0]
                cells = list(first_row.cells)
                for col_idx, cell in enumerate(cells):
                    try:
                        cell_text_upper = cell.text.upper()
                        # Check for LABEL placeholders
                        if ('LABEL' in cell_text_upper or
                            '{{LABEL}}' in cell_text_upper or
                                '{LABEL}' in cell_text_upper):
                            template_cols_with_label.append(col_idx)
                    except (IndexError, AttributeError):
                        continue
            except (IndexError, AttributeError, ValueError):
                for row in working_table.rows:
                    try:
                        cells = list(row.cells)
                        for col_idx, cell in enumerate(cells):
                            try:
                                cell_text_upper = cell.text.upper()
                                if ('LABEL' in cell_text_upper or
                                    '{{LABEL}}' in cell_text_upper or
                                        '{LABEL}' in cell_text_upper) and col_idx not in template_cols_with_label:
                                    template_cols_with_label.append(col_idx)
                            except (IndexError, AttributeError):
                                continue
                        if template_cols_with_label:
                            break
                    except (IndexError, AttributeError, ValueError):
                        continue

        if not template_cols_with_label:
            return doc

        num_cols_available = len(template_cols_with_label)

        sum(len(col_labels) for col_labels in labels_by_column)

        labels_per_table = 0
        for row in first_table.rows:
            try:
                num_cells = len(row.cells) if hasattr(row, 'cells') else 0
                for col_idx in template_cols_with_label:
                    if col_idx < num_cells:
                        try:
                            cell = row.cells[col_idx]
                            if 'LABEL' in cell.text.upper():
                                labels_per_table += 1
                        except (IndexError, AttributeError):
                            continue
            except (IndexError, AttributeError):
                continue

        labels_per_sample_type_per_table = (
            labels_per_table // num_cols_available
            if num_cols_available > 0
            else labels_per_table
        )

        max_labels_per_sample_type = max(len(col_labels) for col_labels in labels_by_column) if labels_by_column else 0
        tables_per_sample_type = max(
            1,
            (max_labels_per_sample_type +
             labels_per_sample_type_per_table -
             1) //
            labels_per_sample_type_per_table) if labels_per_sample_type_per_table > 0 else 1

        tables_needed = (
            max(1, (tables_per_sample_type * num_sample_types +
                    num_cols_available - 1) // num_cols_available)
            if num_cols_available > 0
            else tables_per_sample_type * num_sample_types
        )

        tables_needed = min(tables_needed, 100)

        existing_tables = len(doc.tables)
        tables_to_create = max(0, tables_needed - existing_tables)

        if tables_to_create > 0:
            batch_size = 10
            for batch_start in range(0, tables_to_create, batch_size):
                batch_end = min(batch_start + batch_size, tables_to_create)
                for i in range(batch_start, batch_end):
                    doc.add_page_break()

                    new_tbl = deepcopy(original_template_table_xml)
                    doc._body._body.append(new_tbl)

                    new_table = doc.tables[-1]

                    if original_style_val:
                        try:
                            new_tbl_pr = new_table._tbl.get_or_add_tblPr()

                            has_style_in_xml = False
                            for child in new_tbl_pr:
                                if 'tblStyle' in child.tag:
                                    has_style_in_xml = True
                                    existing_val = child.get(
                                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                    if existing_val != original_style_val:
                                        child.set(
                                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val',
                                            original_style_val)
                                    break

                            if not has_style_in_xml:
                                tbl_style = OxmlElement('w:tblStyle')
                                tbl_style.set(qn('w:val'), original_style_val)
                                new_tbl_pr.append(tbl_style)

                            if original_table_style:
                                new_table.style = original_table_style
                        except Exception:
                            try:
                                if original_table_style:
                                    new_table.style = original_table_style
                            except Exception:
                                pass

        current_sample_type_batch = 0

        filled_table_indices = []

        labels_per_sample_type_per_table = (
            labels_per_table // num_cols_available
            if num_cols_available > 0
            else labels_per_table
        )

        for table_idx, table in enumerate(doc.tables):
            try:
                if len(table.rows) < 1:
                    continue
                if not hasattr(table.rows[0], 'cells') or len(table.rows[0].cells) == 0:
                    continue
            except (IndexError, AttributeError):
                continue

            col_to_sample_type = {}
            if num_sample_types <= num_cols_available:
                for i, template_col_idx in enumerate(template_cols_with_label):
                    label_col_idx = i % num_sample_types
                    if label_index_by_sample_type[label_col_idx] < len(labels_by_column[label_col_idx]):
                        col_to_sample_type[template_col_idx] = label_col_idx
            else:
                for i, template_col_idx in enumerate(template_cols_with_label):
                    sample_type_idx = current_sample_type_batch * num_cols_available + i
                    if sample_type_idx < num_sample_types:
                        if label_index_by_sample_type[sample_type_idx] < len(labels_by_column[sample_type_idx]):
                            col_to_sample_type[template_col_idx] = sample_type_idx

            if not col_to_sample_type:
                current_sample_type_batch += 1
                if num_sample_types <= num_cols_available:
                    for i, template_col_idx in enumerate(template_cols_with_label):
                        label_col_idx = i % num_sample_types
                        if label_index_by_sample_type[label_col_idx] < len(labels_by_column[label_col_idx]):
                            col_to_sample_type[template_col_idx] = label_col_idx
                else:
                    col_to_sample_type = {}
                    for i, template_col_idx in enumerate(template_cols_with_label):
                        sample_type_idx = current_sample_type_batch * num_cols_available + i
                        if sample_type_idx < num_sample_types:
                            if label_index_by_sample_type[sample_type_idx] < len(labels_by_column[sample_type_idx]):
                                col_to_sample_type[template_col_idx] = sample_type_idx

                if not col_to_sample_type:
                    break

            table_was_filled = False

            for row in table.rows:
                try:
                    num_cells = len(row.cells) if hasattr(row, 'cells') else 0
                except (IndexError, AttributeError):
                    continue

                for template_col_idx in col_to_sample_type.keys():
                    if template_col_idx >= num_cells:
                        continue

                    try:
                        cell = row.cells[template_col_idx]
                        cell_text = cell.text.strip()
                        cell_text_upper = cell_text.upper()
                    except (IndexError, AttributeError):
                        continue

                    # Check for LABEL label placeholders
                    if ('LABEL' in cell_text_upper or
                        '{{LABEL}}' in cell_text_upper or
                            '{LABEL}' in cell_text_upper):
                        sample_type_idx = col_to_sample_type[template_col_idx]

                        if label_index_by_sample_type[sample_type_idx] < len(labels_by_column[sample_type_idx]):
                            label_text = labels_by_column[sample_type_idx][label_index_by_sample_type[sample_type_idx]]
                            label_index_by_sample_type[sample_type_idx] += 1
                            table_was_filled = True

                            replaced = False

                            if cell.paragraphs:
                                for paragraph in cell.paragraphs:
                                    if paragraph.runs:
                                        for run in paragraph.runs:
                                            run_text_upper = run.text.upper()
                                            if 'LABEL' in run_text_upper:
                                                fmt_size = run.font.size
                                                fmt_bold = run.font.bold
                                                fmt_name = run.font.name

                                                new_run_text = re.sub(
                                                    r'LABEL', label_text, run.text, flags=re.IGNORECASE)
                                                run.text = new_run_text

                                                if fmt_size:
                                                    run.font.size = fmt_size
                                                if fmt_bold is not None:
                                                    run.font.bold = fmt_bold
                                                if fmt_name:
                                                    run.font.name = fmt_name

                                                replaced = True
                                                break
                                    if replaced:
                                        break

                            if not replaced:
                                new_text = re.sub(r'LABEL', label_text, cell.text, flags=re.IGNORECASE)
                                cell.text = new_text

                                if cell.paragraphs:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            if not run.font.size:
                                                run.font.size = Pt(11)
                                            if run.font.bold is None:
                                                run.font.bold = True

                if all(label_index_by_sample_type[sample_type_idx] >= len(labels_by_column[sample_type_idx])
                       for sample_type_idx in range(num_sample_types)):
                    break

            if table_was_filled:
                filled_table_indices.append(table_idx)

            all_current_batch_filled = True
            for sample_type_idx in col_to_sample_type.values():
                if label_index_by_sample_type[sample_type_idx] < len(labels_by_column[sample_type_idx]):
                    all_current_batch_filled = False
                    break

            if all_current_batch_filled:
                current_sample_type_batch += 1

            if all(label_index_by_sample_type[sample_type_idx] >= len(labels_by_column[sample_type_idx])
                   for sample_type_idx in range(num_sample_types)):
                break

            if table_idx == len(doc.tables) - 1:
                if any(label_index_by_sample_type[sample_type_idx] < len(labels_by_column[sample_type_idx])
                       for sample_type_idx in range(num_sample_types)):
                    try:
                        doc.add_page_break()
                        new_tbl = deepcopy(original_template_table_xml)
                        doc._body._body.append(new_tbl)
                        new_table = doc.tables[-1]
                        if original_table_style:
                            try:
                                new_table.style = original_table_style
                            except Exception:
                                pass
                    except Exception:
                        break

        self._remove_pages_after_last_filled_cell(doc)

        return doc

    def _fill_template_labels(self, doc, labels_by_column):
        """Fill the template with labels by using existing cells - preserves 100% of template formatting"""
        if not doc.tables:
            return doc

        if not labels_by_column:
            return doc

        num_columns = len(labels_by_column)

        label_index_by_column = [0] * num_columns

        filled_tables = []

        for table_idx, table in enumerate(doc.tables):
            if len(table.rows) < 1:
                continue

            template_rows = len(table.rows)
            template_cols = len(table.rows[0].cells) if table.rows else 0

            if num_columns < template_cols:
                for row_idx in range(template_rows):
                    row = table.rows[row_idx]

                    for template_col_idx in range(template_cols):
                        if template_col_idx >= len(row.cells):
                            break

                        cell = row.cells[template_col_idx]
                        label_col_idx = template_col_idx % num_columns
                        column_labels = labels_by_column[label_col_idx]

                        if label_index_by_column[label_col_idx] < len(column_labels):
                            label_text = column_labels[label_index_by_column[label_col_idx]]
                            label_index_by_column[label_col_idx] += 1
                            table_was_filled = True
                        else:
                            label_text = ''

                        self._replace_cell_text_preserving_format(cell, label_text)
            else:
                if num_columns > template_cols:
                    continue

                table_was_filled = False

                for row_idx in range(template_rows):
                    row = table.rows[row_idx]

                    for col_idx in range(num_columns):
                        if col_idx >= len(row.cells):
                            break

                        cell = row.cells[col_idx]
                        column_labels = labels_by_column[col_idx]

                        if label_index_by_column[col_idx] < len(column_labels):
                            label_text = column_labels[label_index_by_column[col_idx]]
                            label_index_by_column[col_idx] += 1
                            table_was_filled = True
                        else:
                            label_text = ''

                        self._replace_cell_text_preserving_format(cell, label_text)

            if table_was_filled:
                filled_tables.append(table_idx)

            if all(label_index_by_column[col_idx] >= len(labels_by_column[col_idx])
                   for col_idx in range(num_columns)):
                break

        self._remove_pages_after_last_filled_cell(doc)

        return doc

    def _shade_cell(self, cell, fill_hex):
        """
        Apply a solid background fill to a table cell using the w:shd XML element.
        """
        try:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()

            for existing in tc_pr.findall(qn('w:shd')):
                tc_pr.remove(existing)

            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill_hex.upper().lstrip('#'))
            tc_pr.append(shd)
        except Exception:
            pass

    def _string_contains_label_placeholder(self, text):
        """True if plain text is a label placeholder"""
        try:
            u = (text or '').upper()
        except (AttributeError, TypeError):
            return False
        return 'LABEL' in u or '{{LABEL}}' in u or '{LABEL}' in u

    def _cell_contains_label_placeholder(self, cell):
        """True if this cell is intended as a label slot"""
        try:
            return self._string_contains_label_placeholder(cell.text)
        except (AttributeError, TypeError):
            return False

    def _tc_preferred_width_twips(self, tc):
        """Return w:tc preferred width in twips if set; None if missing or invalid."""
        try:
            tc_pr = tc.find(qn('w:tcPr'))
            if tc_pr is None:
                return None
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                return None
            w = tc_w.get(qn('w:w'))
            if w is None:
                return None
            return int(w)
        except (TypeError, ValueError, AttributeError):
            return None

    def _collect_placeholder_cells_xml(self, table):
        tbl = table._tbl
        cells_out = []
        # Only direct child w:tr of the table. tbl.iter(w:tr) would include rows of nested tables
        # inside a cell, corrupting slot order and count.
        for tr in tbl:
            if tr.tag != qn('w:tr'):
                continue
            row = _Row(tr, table)
            # Only direct child w:tc per row. tr.iter(w:tc) would include nested-table cells and
            # inflate the slot list, so two specs can target the same visual label and one transect
            # disappears (e.g. 48 codes → 47 on single-per-transect types).
            for tc in tr:
                if tc.tag != qn('w:tc'):
                    continue
                texts = []
                for wt in tc.iter(qn('w:t')):
                    if wt.text:
                        texts.append(wt.text)
                blob = ''.join(texts)
                if self._string_contains_label_placeholder(blob):
                    tw = self._tc_preferred_width_twips(tc)
                    if (
                        tw is not None
                        and tw < self._MIN_LABEL_SLOT_CELL_WIDTH_TWIPS
                    ):
                        continue
                    # Same physical tc should not be listed twice (would double-write and lose labels).
                    if any(c._tc is tc for c in cells_out):
                        continue
                    cells_out.append(_Cell(tc, row))
        return cells_out

    def _count_placeholder_tcs_in_table_xml(self, table):
        return len(self._collect_placeholder_cells_xml(table))

    def _strip_trailing_empty_paragraphs_before_sectpr(self, doc):
        body = doc._element.body
        while True:
            elems = list(body)
            try:
                si = next(i for i, el in enumerate(elems) if el.tag == qn('w:sectPr'))
            except StopIteration:
                return
            if si == 0:
                return
            prev = elems[si - 1]
            if prev.tag != qn('w:p'):
                return
            if ''.join(prev.itertext()).strip() != '':
                return
            body.remove(prev)

    def _select_label_template_table(self, doc):

        best_table = None
        best_count = 0
        best_rows = 0
        for table in doc.tables:
            try:
                count = self._count_placeholder_tcs_in_table_xml(table)
                row_count = len(table.rows)
                if count > best_count or (count == best_count and row_count > best_rows):
                    best_count = count
                    best_rows = row_count
                    best_table = table
            except (ValueError, AttributeError, IndexError, TypeError):
                continue
        return best_table, best_count

    def _fill_template_by_groups(self, doc, groups_of_label_specs, group_colors=None):
        """
        Fill the template with groups of labels – one group per sample type / label category.
        """
        ADDRESS_BLOCK = "Ecdysis Foundation\n46958 188th St\nEstelline, SD  57234"
        if not doc.tables or not groups_of_label_specs:
            return doc

        self._strip_trailing_empty_paragraphs_before_sectpr(doc)

        template_table, _label_slot_count = self._select_label_template_table(doc)

        if template_table is None:
            flat = []
            for g in groups_of_label_specs:
                for spec in g:
                    flat.append(spec.get('text', ''))
            return self._create_tables_for_labels(doc, [flat])

        try:
            template_table_idx = next(
                i for i, tbl in enumerate(doc.tables) if tbl._tbl is template_table._tbl
            )
        except StopIteration:
            template_table_idx = 0

        template_table_xml = deepcopy(template_table._tbl)
        original_table_style = None
        original_style_val = None
        try:
            original_table_style = template_table.style
            original_tbl_pr = template_table._tbl.tblPr
            if original_tbl_pr is not None:
                for child in original_tbl_pr:
                    if 'tblStyle' in child.tag:
                        original_style_val = child.get(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                        break
        except Exception:
            pass

        try:
            template_label_cells = self._collect_placeholder_cells_xml(template_table)
        except (ValueError, AttributeError, TypeError):
            template_label_cells = []

        if not template_label_cells:
            return doc

        def _add_template_table():
            """Append a page break + fresh template-table copy to the document."""
            doc.add_page_break()
            new_tbl = deepcopy(template_table_xml)
            doc._body._body.append(new_tbl)
            new_table = doc.tables[-1]
            if original_style_val:
                try:
                    new_tbl_pr = new_table._tbl.get_or_add_tblPr()
                    has_style = any('tblStyle' in child.tag for child in new_tbl_pr)
                    if has_style:
                        for child in new_tbl_pr:
                            if 'tblStyle' in child.tag:
                                child.set(
                                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val',
                                    original_style_val)
                                break
                    else:
                        tbl_style = OxmlElement('w:tblStyle')
                        tbl_style.set(qn('w:val'), original_style_val)
                        new_tbl_pr.append(tbl_style)
                    if original_table_style:
                        new_table.style = original_table_style
                except Exception:
                    try:
                        if original_table_style:
                            new_table.style = original_table_style
                    except Exception:
                        pass
            return new_table

        def _write_cell_from_spec(cell, spec):
            """write text to a cell"""
            try:
                text = (spec or {}).get('text', '')
                font_size = (spec or {}).get('font_size')
                bold_all = bool((spec or {}).get('bold_all', False))
                bold_first_line = bool((spec or {}).get('bold_first_line', False))
                alignment = (spec or {}).get('alignment', 'center')

                cell.text = ""
                paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                if alignment == 'left':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif alignment == 'right':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                lines = str(text).splitlines() if text is not None else [""]
                for i, line in enumerate(lines):
                    run = paragraph.add_run(line)
                    if font_size:
                        run.font.size = Pt(int(font_size))
                    run.font.name = 'Arial'
                    run.font.bold = True if (bold_all or (bold_first_line and i == 0)) else False
                    if i < len(lines) - 1:
                        run.add_break()
            except Exception:
                try:
                    cell.text = (spec or {}).get('text', '')
                except Exception:
                    pass

            try:
                fill_hex = (spec or {}).get('cell_fill_hex')
                if fill_hex:
                    self._shade_cell(cell, fill_hex)
            except Exception:
                pass

        def _fill_placeholder_cells_for_table(
            table,
            specs_slice,
            fill_color=None,
            fill_remaining_spec=None,
            placeholder_cells=None,
            fill_log_context=None,
        ):
            """Fill one sheet's worth of label cells (same layout as template table)."""
            filled_specs = 0
            cells = placeholder_cells
            if cells is None:
                try:
                    cells = self._collect_placeholder_cells_xml(table)
                except (ValueError, AttributeError, TypeError):
                    return 0
            if not cells:
                return 0
            debug_cells = _label_fill_debug_enabled()
            gix = (fill_log_context or {}).get('group_idx')
            spec0 = (fill_log_context or {}).get('spec_index_start')
            for idx, cell in enumerate(cells):
                if idx < len(specs_slice):
                    spec = dict(specs_slice[idx] or {})
                    if fill_color and not spec.get('cell_fill_hex'):
                        spec['cell_fill_hex'] = fill_color
                    _write_cell_from_spec(cell, spec)
                    filled_specs += 1
                    if debug_cells and gix is not None and spec0 is not None:
                        preview = (spec.get('text') or '')[:120].replace('\n', '|')
                        logger.debug(
                            'labels.fill.write group=%s abs_spec=%s cell_slot=%s/%s preview=%r',
                            gix,
                            spec0 + idx,
                            idx,
                            len(cells) - 1,
                            preview,
                        )
                elif fill_remaining_spec is not None:
                    spec = dict(fill_remaining_spec or {})
                    if fill_color and not spec.get('cell_fill_hex'):
                        spec['cell_fill_hex'] = fill_color
                    _write_cell_from_spec(cell, spec)
                    if debug_cells and gix is not None:
                        logger.debug(
                            'labels.fill.address group=%s cell_slot=%s/%s',
                            gix,
                            idx,
                            len(cells) - 1,
                        )
            return filled_specs

        first_group = True
        for group_idx, group_specs in enumerate(groups_of_label_specs):
            if not group_specs:
                continue

            fill_color = (
                group_colors[group_idx]
                if group_colors and group_idx < len(group_colors)
                else None
            )

            if not first_group:
                _add_template_table()

            if _label_fill_info_enabled():
                logger.info(
                    'labels.fill.group_start group=%s specs=%s project=%s category=%s',
                    group_idx,
                    len(group_specs),
                    self.project_type,
                    self.label_category,
                )

            label_idx = 0
            while label_idx < len(group_specs):
                # First page of the first sample-type group must use the template's label table,
                # not doc.tables[-1] (templates often have extra tables after the label grid).
                if first_group and label_idx == 0:
                    current_table = doc.tables[template_table_idx]
                else:
                    current_table = doc.tables[-1]
                try:
                    page_cells = self._collect_placeholder_cells_xml(current_table)
                except (ValueError, AttributeError, TypeError):
                    page_cells = []
                page_capacity = len(page_cells)
                if page_capacity == 0:
                    logger.error(
                        'labels.fill.zero_capacity group=%s label_idx=%s of %s doc_tables=%s',
                        group_idx,
                        label_idx,
                        len(group_specs),
                        len(doc.tables),
                    )
                    break
                remaining = len(group_specs) - label_idx
                chunk = min(page_capacity, remaining)
                specs_slice = group_specs[label_idx:label_idx + chunk]

                is_last_table_for_group = (label_idx + chunk) >= len(group_specs)
                skip_address_fill = (
                    self.project_type == 'ignite' and self.label_category == 'outer'
                )
                _addr_fs = 12 if self.project_type == 'ignite' else 16
                fill_remaining_spec = (
                    {'text': ADDRESS_BLOCK, 'font_size': _addr_fs, 'bold_all': True, 'bold_first_line': False}
                    if is_last_table_for_group and not skip_address_fill
                    else None
                )

                try:
                    cur_tbl_i = next(
                        i for i, t in enumerate(doc.tables) if t._tbl is current_table._tbl
                    )
                except StopIteration:
                    cur_tbl_i = -1

                filled_specs = _fill_placeholder_cells_for_table(
                    current_table,
                    specs_slice,
                    fill_color=fill_color,
                    fill_remaining_spec=fill_remaining_spec,
                    placeholder_cells=page_cells,
                    fill_log_context={
                        'group_idx': group_idx,
                        'spec_index_start': label_idx,
                    },
                )
                if _label_fill_info_enabled():
                    logger.info(
                        'labels.fill.page group=%s table_idx=%s label_range=[%s,%s) chunk=%s '
                        'capacity=%s last_table_for_group=%s filled_specs=%s',
                        group_idx,
                        cur_tbl_i,
                        label_idx,
                        label_idx + chunk,
                        chunk,
                        page_capacity,
                        is_last_table_for_group,
                        filled_specs,
                    )
                if filled_specs != len(specs_slice):
                    logger.warning(
                        'labels.fill.filled_vs_slice group=%s filled_specs=%s slice_len=%s',
                        group_idx,
                        filled_specs,
                        len(specs_slice),
                    )
                if filled_specs == 0 and not fill_remaining_spec:
                    logger.error(
                        'labels.fill.zero_fill group=%s label_idx=%s chunk=%s capacity=%s',
                        group_idx,
                        label_idx,
                        chunk,
                        page_capacity,
                    )
                    break
                # Advance by specs actually written from specs_slice (must match len(specs_slice)
                # when every cell received a spec or trailing address; avoids skipping if slice
                # length ever disagreed with the cell list used for filling).
                label_idx += filled_specs
                if label_idx < len(group_specs):
                    _add_template_table()

            if label_idx != len(group_specs):
                logger.error(
                    'labels.fill.incomplete_group group=%s wrote_specs=%s expected=%s',
                    group_idx,
                    label_idx,
                    len(group_specs),
                )

            first_group = False

        self._remove_pages_after_last_filled_cell(doc)
        self._strip_trailing_empty_paragraphs_before_sectpr(doc)
        return doc

    def generate_quick_labels_avalanche(self, num_transects):
        """Generate all labels for Avalanche project. complete set per transect"""

        transect_codes = self.generate_unique_sample_codes(num_transects)
        self.save_sample_codes(transect_codes)

        logger.info(
            'labels.quick.transects requested=%s generated=%s cluster=%s year=%s fill_info=%s fill_debug=%s',
            num_transects,
            len(transect_codes),
            self.cluster_number,
            self.year,
            _label_fill_info_enabled(),
            _label_fill_debug_enabled(),
        )
        if len(transect_codes) != num_transects:
            msg = (
                f'Expected {num_transects} unique Avalanche transect codes, '
                f'but generate_unique_sample_codes returned {len(transect_codes)}.'
            )
            logger.error('labels.quick.count_mismatch %s', msg)
            raise ValueError(msg)
        return self._build_quick_avalanche_docx(transect_codes, expected_len=num_transects)

    def regenerate_quick_avalanche_inner_docx(self, transect_codes):
        """
        Rebuild the inner quick-label Word file from existing codes only (no new SampleCode rows).

        Used when an earlier document was produced with buggy fill logic but the same codes
        should appear in the file.
        """
        codes = [str(c).strip() for c in (transect_codes or []) if str(c).strip()]
        if not codes:
            raise ValueError('No transect codes supplied; cannot regenerate the label document.')
        self.generated_codes = codes
        expected = self.labels_per_type or len(codes)
        if len(codes) != expected:
            raise ValueError(
                f'This batch is configured for {expected} transect codes but '
                f'{len(codes)} non-empty code(s) were provided.'
            )
        logger.info(
            'labels.quick.regenerate codes=%s cluster=%s year=%s fill_info=%s fill_debug=%s',
            len(codes),
            self.cluster_number,
            self.year,
            _label_fill_info_enabled(),
            _label_fill_debug_enabled(),
        )
        return self._build_quick_avalanche_docx(codes, expected_len=expected)

    def _build_quick_avalanche_docx(self, transect_codes, *, expected_len):
        """Shared body for quick inner Avalanche: build groups, fill template, return buffer."""
        if len(transect_codes) != expected_len:
            msg = (
                f'Internal error: expected {expected_len} transect codes in document build, '
                f'got {len(transect_codes)}.'
            )
            logger.error('labels.quick.build_len_mismatch %s', msg)
            raise ValueError(msg)
        # consider excluded sample types passed in from the form
        sample_type_codes = list(self.sample_types) if self.sample_types else [code for code, _ in SAMPLE_TYPES]
        ordered_type_codes = self._avalanche_inner_sample_types_for_fill_order(sample_type_codes)
        ordered_type_codes = [c for c in ordered_type_codes if c != 'soil_core_0_60cm']

        doc = self._load_template()

        groups_of_label_specs = []
        total_labels = 0

        for sample_type_code in ordered_type_codes:
            sample_type_display = self.get_sample_type_display(sample_type_code)
            labels_per_transect = self._avalanche_inner_labels_per_transect(sample_type_code)
            group = []
            for transect_code in transect_codes:
                label_text = self.create_label_text(sample_type_display, transect_code)
                for _ in range(labels_per_transect):
                    group.append({
                        'text': label_text,
                        'font_size': 16,
                        'bold_all': False,
                        'bold_first_line': False,
                    })
                    total_labels += 1
            groups_of_label_specs.append(group)

        if 'soil_core_0_60cm' in sample_type_codes:
            special_group = []
            for transect_code in transect_codes:
                special_text = (
                    "Soil Core 0-60cm\n"
                    f"Cluster {self.cluster_number} – {self.year}\n"
                    f"{transect_code}"
                )
                special_group.append({
                    'text': special_text,
                    'font_size': 16,
                    'bold_all': False,
                    'bold_first_line': False,
                })
                total_labels += 1
            groups_of_label_specs.append(special_group)

        if doc.tables:
            doc = self._fill_template_by_groups(doc, groups_of_label_specs)
        else:
            flat = []
            for g in groups_of_label_specs:
                flat.extend([spec.get('text', '') for spec in g])
            doc = self._create_tables_for_labels(doc, [flat])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer, total_labels

    def generate_quick_labels_ignite(self, num_sites):
        """Generate all inner labels for Ignite project"""
        site_codes, start_num = self.generate_unique_site_codes(num_sites)

        saved_codes = []
        for code in site_codes:
            tc = SampleCode.objects.create(
                code=code,
                project_type='ignite',
                cluster_number=self.cluster_number,
                year=self.year,
                site_code_numeric=int(code),
                created_by=self.created_by,
                is_active=True
            )
            saved_codes.append(code)

            # Create 4 transects for each site
            for t_num in range(1, 5):
                SiteTransect.objects.create(
                    site_code=tc,
                    transect_number=t_num,
                    is_active=True
                )

        self.generated_codes = saved_codes

        doc = self._load_template()
        total_labels = 0

        groups_of_label_specs = []
        ordered_types = self._ignite_inner_sample_types_for_fill_order(self.sample_types)
        for sample_type_code in ordered_types:
            sample_type_display = self.get_sample_type_display(sample_type_code)
            group = []
            for site_code in site_codes:
                for t_num in range(1, 5):
                    label_text = self.create_ignite_inner_label(
                        sample_type_display, site_code, t_num
                    )
                    group.append({
                        'text': label_text,
                        'font_size': 12,
                        'bold_all': False,
                        'bold_first_line': False,
                    })
                    total_labels += 1
            groups_of_label_specs.append(group)

        if doc.tables:
            doc = self._fill_template_by_groups(doc, groups_of_label_specs)
        else:
            flat = []
            for g in groups_of_label_specs:
                flat.extend([spec.get('text', '') for spec in g])
            doc = self._create_tables_for_labels(doc, [flat])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer, total_labels

    def regenerate_quick_ignite_inner_docx(self, site_codes, *, expected_num_sites: int | None = None):
        """
        Rebuild Ignite inner labels from existing site codes only (no new SampleCode / SiteTransect rows).

        ``transect_codes_generated`` on Ignite inner batches stores site codes (one per farm/site).
        """
        codes = [str(c).strip() for c in (site_codes or []) if str(c).strip()]
        if not codes:
            raise ValueError('No site codes supplied; cannot regenerate Ignite inner labels.')
        expected = expected_num_sites if expected_num_sites is not None else len(codes)
        if len(codes) != expected:
            raise ValueError(
                f'Expected {expected} Ignite site code(s) for this batch but received '
                f'{len(codes)} non-empty value(s).'
            )
        self.generated_codes = codes
        logger.info(
            'labels.quick.regenerate_ignite_inner sites=%s cluster=%s year=%s fill_info=%s fill_debug=%s',
            len(codes),
            self.cluster_number,
            self.year,
            _label_fill_info_enabled(),
            _label_fill_debug_enabled(),
        )

        doc = self._load_template()
        total_labels = 0

        groups_of_label_specs = []
        ordered_types = self._ignite_inner_sample_types_for_fill_order(self.sample_types)
        for sample_type_code in ordered_types:
            sample_type_display = self.get_sample_type_display(sample_type_code)
            group = []
            for site_code in codes:
                for t_num in range(1, 5):
                    label_text = self.create_ignite_inner_label(
                        sample_type_display, site_code, t_num
                    )
                    group.append({
                        'text': label_text,
                        'font_size': 12,
                        'bold_all': False,
                        'bold_first_line': False,
                    })
                    total_labels += 1
            groups_of_label_specs.append(group)

        if doc.tables:
            doc = self._fill_template_by_groups(doc, groups_of_label_specs)
        else:
            flat = []
            for g in groups_of_label_specs:
                flat.extend([spec.get('text', '') for spec in g])
            doc = self._create_tables_for_labels(doc, [flat])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer, total_labels

    def _fill_ignite_labels_sequential(self, doc, all_labels):
        """Fill Ignite labels"""
        if not doc.tables or not all_labels:
            return doc

        template_table = None
        for idx, table in enumerate(doc.tables):
            try:
                rows = list(table.rows)
                if not rows:
                    continue
                cells = list(rows[0].cells)
                has_label = any('LABEL' in cell.text.upper() for cell in cells)
                if has_label:
                    template_table = table
                    print(f"DEBUG: Found template table at index {idx}")
                    break
            except (ValueError, AttributeError):
                continue

        if not template_table:
            print("DEBUG: No template table found!")
            return doc

        labels_per_page = 0
        label_cells_per_row = []

        try:
            for row in template_table.rows:
                cells = list(row.cells)
                row_label_indices = []
                for cell_idx, cell in enumerate(cells):
                    if 'LABEL' in cell.text.upper():
                        row_label_indices.append(cell_idx)
                        labels_per_page += 1
                label_cells_per_row.append(row_label_indices)
        except (ValueError, AttributeError) as e:
            print(f"DEBUG: Error counting labels: {e}")
            return doc

        pages_needed = (len(all_labels) + labels_per_page - 1) // labels_per_page

        template_table_xml = deepcopy(template_table._tbl)

        for page_num in range(1, pages_needed):
            doc.add_page_break()
            new_tbl = deepcopy(template_table_xml)
            doc._body._body.append(new_tbl)
            print(f"DEBUG: Added page {page_num + 1}")

        label_idx = 0

        fillable_tables = []

        if not label_cells_per_row:
            print("DEBUG: ERROR - label_cells_per_row is empty!")
            return doc

        max_cell_idx = max(max(row_indices)
                           for row_indices in label_cells_per_row if row_indices) if label_cells_per_row else 0

        for table_num, table in enumerate(doc.tables):
            try:
                rows = list(table.rows)
                if rows and len(rows) >= len(label_cells_per_row):
                    test_row = rows[0]
                    test_cells = list(test_row.cells)
                    if len(test_cells) > max_cell_idx:
                        fillable_tables.append((table_num, table))
            except (ValueError, AttributeError, IndexError) as e:
                print(f"DEBUG: Skipping table {table_num} (error: {e})")
                continue

        print(f"DEBUG: Found {len(fillable_tables)} fillable tables")

        for table_num, table in fillable_tables:
            if label_idx >= len(all_labels):
                break

            try:
                rows = list(table.rows)
                for row_idx, row in enumerate(rows):
                    if label_idx >= len(all_labels):
                        break

                    if row_idx >= len(label_cells_per_row) or row_idx < 0:
                        continue

                    row_label_indices = label_cells_per_row[row_idx]
                    if not row_label_indices:
                        continue

                    try:
                        cells = list(row.cells)
                        for cell_idx in row_label_indices:
                            if label_idx >= len(all_labels):
                                break

                            if cell_idx < len(cells):
                                cell = cells[cell_idx]
                                label_text = all_labels[label_idx]

                                if cell.paragraphs:
                                    cell.paragraphs[0].clear()
                                    run = cell.paragraphs[0].add_run(label_text)
                                    run.font.size = Pt(9)

                                label_idx += 1
                            else:
                                print(
                                    f"DEBUG: Warning: cell_idx {cell_idx} >= len(cells) {
                                        len(cells)} in table {table_num} row {row_idx}")
                    except (ValueError, AttributeError, IndexError) as e:
                        print(f"DEBUG: Error accessing row {row_idx} in table {table_num}: {e}")
                        continue

            except (ValueError, AttributeError, IndexError) as e:
                print(f"DEBUG: Error filling table {table_num}: {e}")
                continue

        print(f"DEBUG: Filled {label_idx} labels")
        return doc

    def create_ignite_inner_label(self, sample_type, site_code, transect_num):
        """Format: Cluster XX – YYYY / Sample Type TX / Site: XXXX"""
        return f"Cluster {self.cluster_number} – {self.year}\n{sample_type} T{transect_num}\nSite: {site_code}"

    @staticmethod
    def _ignite_inner_sample_types_for_fill_order(sample_type_codes):
        seq = list(sample_type_codes)
        rest = [c for c in seq if c not in ('plant_dna', 'forage')]
        tail = []
        if 'forage' in seq:
            tail.append('forage')
        if 'plant_dna' in seq:
            tail.append('plant_dna')
        return rest + tail

    @staticmethod
    def _avalanche_inner_sample_types_for_fill_order(sample_type_codes):
        seq = list(sample_type_codes)
        return [c for c in seq if c != 'plant_dna'] + (['plant_dna'] if 'plant_dna' in seq else [])

    @staticmethod
    def _avalanche_inner_labels_per_transect(sample_type_code):
        if sample_type_code == 'plant_dna':
            return 2
        if sample_type_code == 'forage':
            return 2
        if sample_type_code == 'yield_sample':
            return 2
        return 1

    def create_ignite_outer_label(self, sample_type, site_code):
        """Format: Cluster XX – YYYY / Sample Type / Site: XXXX"""
        return f"Cluster {self.cluster_number} – {self.year}\n{sample_type}\nSite: {site_code}"

    def create_ignite_crop_type_variety_outer_label(self, site_code):
        """Format: Cluster XX – YYYY / Type / Variety-Breed blanks / Site: XXXX"""
        return (
            f"Cluster {self.cluster_number} – {self.year}\n"
            "Type ________\n"
            "Variety/Breed ________\n"
            f"Site: {site_code}"
        )

    def generate_outer_labels_ignite(self, site_codes):
        """Generate outer labels for Ignite – one sample type per page."""
        doc = self._load_template()
        total_labels = 0

        # One group per sample type; each type gets its own page(s).
        groups_of_label_specs = []
        for sample_type_code in self.sample_types:
            group = []
            for site_code in site_codes:
                if sample_type_code == 'crop_type_variety':
                    label_text = self.create_ignite_crop_type_variety_outer_label(site_code)
                    group.append({
                        'text': label_text,
                        'font_size': 9,
                        'alignment': 'left',
                        'bold_all': False,
                        'bold_first_line': False,
                    })
                else:
                    sample_type_display = self.get_sample_type_display(sample_type_code)
                    label_text = self.create_ignite_outer_label(sample_type_display, site_code)
                    group.append({
                        'text': label_text,
                        'font_size': 12,
                        'bold_all': False,
                        'bold_first_line': False,
                    })
                total_labels += 1
            groups_of_label_specs.append(group)

        if doc.tables:
            doc = self._fill_template_by_groups(doc, groups_of_label_specs)
        else:
            flat = []
            for g in groups_of_label_specs:
                flat.extend([spec.get('text', '') for spec in g])
            doc = self._create_tables_for_labels(doc, [flat])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer, total_labels

    def generate_outer_labels_avalanche(self, transect_codes):
        """Generate outer labels for Avalanche – one label category per page."""
        doc = self._load_template()
        total_labels = 0

        kit_labels = []
        room_temp_labels = []
        refrigerated_labels = []
        booklet_labels = []

        for transect_code in transect_codes:
            kit_text = (
                "Avalanche Sampling Kit\n"
                "Ecdysis Foundation\n"
                f"Cluster {self.cluster_number} – {self.year}\n"
                f"{transect_code}\n"
                "Grower Name______________\n"
                "Date_____________"
            )
            kit_labels.append({
                'text': kit_text,
                'font_size': 16,
                'bold_all': True,
                'bold_first_line': False,
            })

            room_temp_text = (
                "Avalanche Samples - Room Temp\n"
                f"Cluster {self.cluster_number} – {self.year}\n"
                f"{transect_code}\n"
                "Grower Name______________\n"
                "Date_____________\n" +
                "\n".join(AVALANCHE_ROOM_TEMP_SAMPLES)
            )
            room_temp_labels.append({
                'text': room_temp_text,
                'font_size': 9,
                'bold_all': False,
                'bold_first_line': True,
            })

            refrigerated_text = (
                "Avalanche Samples - REFRIGERATED\n"
                f"Cluster {self.cluster_number} – {self.year}\n"
                f"{transect_code}\n"
                "Grower Name______________\n"
                "Date_____________\n" +
                "\n".join(AVALANCHE_REFRIGERATED_SAMPLES)
            )
            refrigerated_labels.append({
                'text': refrigerated_text,
                'font_size': 11,
                'bold_all': False,
                'bold_first_line': True,
                'cell_fill_hex': 'CCE5FF',
            })

            booklet_text = (
                "Sampling Booklet\n"
                f"Cluster {self.cluster_number} – {self.year}\n"
                f"{transect_code}\n\n"
                "Grower Name______________"
            )
            booklet_labels.append({
                'text': booklet_text,
                'font_size': 16,
                'bold_all': False,
                'bold_first_line': False,
            })
            total_labels += 4

        groups_of_label_specs = [kit_labels, room_temp_labels, refrigerated_labels, booklet_labels]

        if doc.tables:
            doc = self._fill_template_by_groups(doc, groups_of_label_specs)
        else:
            flat = []
            for g in groups_of_label_specs:
                flat.extend([spec.get('text', '') for spec in g])
            doc = self._create_tables_for_labels(doc, [flat])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer, total_labels

    def _create_tables_for_labels(self, doc, labels_by_column):
        """Create tables from scratch for labels"""
        max_rows = max(len(col) for col in labels_by_column) if labels_by_column else 0
        table_count = (len(labels_by_column) + 2) // 3

        for table_idx in range(table_count):
            start_col = table_idx * 3
            end_col = min(start_col + 3, len(labels_by_column))
            cols_in_table = end_col - start_col

            if table_idx > 0:
                doc.add_page_break()

            table = doc.add_table(rows=max_rows, cols=cols_in_table)
            table.style = 'Table Grid'

            for col_idx in range(cols_in_table):
                sample_idx = start_col + col_idx
                column_labels = labels_by_column[sample_idx]

                for row_idx in range(max_rows):
                    cell = table.rows[row_idx].cells[col_idx]

                    if row_idx < len(column_labels):
                        cell.text = column_labels[row_idx]
                        cell.width = Inches(2.3)

                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(11)
                                run.font.bold = True

            for row in table.rows:
                row.height = Inches(0.85)

        return doc

    def generate_document_with_table(self):
        """Generate Word document with labels"""
        total_labels = len(self.sample_types) * self.labels_per_type

        transect_codes = self.generate_unique_sample_codes(total_labels)
        self.save_sample_codes(transect_codes)

        doc = self._load_template()

        code_index = 0
        groups_of_label_specs = []

        for sample_type_code in self.sample_types:
            sample_type_display = self.get_sample_type_display(sample_type_code)
            group = []

            for _ in range(self.labels_per_type):
                transect_code = transect_codes[code_index]
                label_text = self.create_label_text(sample_type_display, transect_code)
                group.append({
                    'text': label_text,
                    'font_size': (
                        12
                        if self.project_type == 'ignite'
                        else (16 if self.project_type == 'avalanche' else 11)
                    ),
                    'bold_all': False,
                    'bold_first_line': False,
                })
                code_index += 1

            groups_of_label_specs.append(group)

        if doc.tables:
            doc = self._fill_template_by_groups(doc, groups_of_label_specs)
        else:
            flat = []
            for g in groups_of_label_specs:
                flat.extend([spec.get('text', '') for spec in g])
            doc = self._create_tables_for_labels(doc, [flat])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer, total_labels
